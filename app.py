import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from docx import Document as DocxReader
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import shutil
import time # Import the time module for streaming effect
from langchain.docstore.document import Document

# Load environment variables (like GOOGLE_API_KEY) from a .env file
load_dotenv()

# Configure Google Generative AI with the API key
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

# --- Helper Functions for PDF Processing ---

from docx import Document as DocxReader  # Required to handle .docx files
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
import streamlit as st

def get_pdf_documents(uploaded_docs):
    """
    Extracts text from uploaded .pdf and .docx documents safely, preserving metadata.
    """
    documents = []

    for doc in uploaded_docs:
        file_name = doc.name.lower()

        try:
            if file_name.endswith(".pdf"):
                # Use PdfReader only on PDFs
                pdf_reader = PdfReader(doc)
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        documents.append(Document(
                            page_content=page_text,
                            metadata={"page": i + 1, "source": doc.name}
                        ))

            elif file_name.endswith(".docx"):
                # Use python-docx only on Word files
                docx_reader = DocxReader(doc)
                full_text = []
                for para in docx_reader.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                full_doc_text = "\n".join(full_text)
                documents.append(Document(
                    page_content=full_doc_text,
                    metadata={"page": 1, "source": doc.name}
                ))

            else:
                st.warning(f"❌ Unsupported file type: {doc.name}. Only PDF and DOCX are supported.")

        except Exception as e:
            st.warning(f"⚠️ Could not read '{doc.name}': {e}")

    return documents


def get_text_chunks_from_documents(documents):
    """
    Splits LangChain Document objects into smaller, overlapping chunks,
    preserving metadata.
    """
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_documents(documents)
    return chunks

def get_vector_store(text_chunks):
    """
    Creates and saves a FAISS vector store from text chunks (LangChain Document objects).
    """
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_documents(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")
    st.session_state.vector_store_loaded = True

def get_conversational_chain_components():
    """
    Defines the components (model and prompt template) for PDF Q&A.
    """
    prompt_template = """
    You are an expert document assistant. Answer the question as detailed as possible *from the provided context only*.
    **It is CRUCIAL that you format your answer using Markdown.**
    **Use headings (e.g., ## Main Topic, ### Sub-section) and bold important keywords or phrases (e.g., **key detail**) to make the answer highly readable and structured, just like a well-formatted document.**
    Ensure all relevant details from the context are included.
    **Always include the page number(s) from which the information was extracted in your answer, e.g., (Page X, Y). If information spans multiple pages, list all relevant pages.**

    **If the answer is NOT available in the provided context, you MUST ONLY respond with: "Answer is not available in the context." Do NOT invent information or provide a wrong answer.**

    Context:\n {context}\n
    Question: \n{question}\n

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return model, prompt

# --- Helper Functions for General Chatbot ---

general_chat_model = genai.GenerativeModel('gemini-2.0-flash')

def translate_role_for_streamlit(user_role):
    if user_role == "model":
        return "assistant"
    else:
        return user_role

# --- Main Streamlit Application ---

def main():
    st.set_page_config(page_title="Chat PDF", page_icon="📄", layout="wide")

         # Developer badge in top-right corner
    st.markdown("""
        <style>
        .top-right-badge {
            position: fixed;
            top: 15px;
            right: 25px;
            background-color: #f0f2f6;
            padding: 8px 15px;
            border-radius: 10px;
            font-weight: bold;
            box-shadow: 0 2px 6px rgba(0,0,0,0.1);
            z-index: 9999;
        }
        </style>
<div class="top-right-badge">
    <span style='color: black;'>Developed by</span> <a href='https://anasaamer.github.io/Protfolio-Website/' target='_blank' style='text-decoration: none; color: #4B8BBE;'>Anas Aamer</a>
</div>
""", unsafe_allow_html=True)

    # === MAIN PAGE ===
    st.markdown("<h1 style='text-align: center; color:#4B8BBE;'>📚 ChatPDF</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Upload multiple PDF files and chat with them intelligently using AI.</p>", unsafe_allow_html=True)

    st.divider()

    # Add "Hello, WhatsApp" centered below the previous paragraph
    st.markdown("<h2 style='text-align: center; color: #4B8BBE; font-weight: bold; padding-left: 50px;'>Hello, WhatsApp</h2>", unsafe_allow_html=True)


    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "vector_store_loaded" not in st.session_state:
        st.session_state.vector_store_loaded = False
    if "chat_session" not in st.session_state:
        st.session_state.chat_session = general_chat_model.start_chat(history=[])
    if "current_chat_mode" not in st.session_state:
        st.session_state.current_chat_mode = "general"

    st.markdown("""
    <style>
    /* General chat message styling */
    .stChatMessage {
        border-radius: 15px;
        padding: 10px 15px;
        margin-bottom: 10px;
        max-width: 80%;
    }

    /* User message styling */
    .stChatMessage[data-testid="stChatMessage-user"] {
        background-color: #e0e0e0;
        align-self: flex-end;
        margin-left: auto;
    }

    /* Assistant message styling */
    .stChatMessage[data-testid="stChatMessage-assistant"] {
        background-color: #f0f2f6;
        align-self: flex-start;
        margin-right: auto;
    }

    /* Hide the Streamlit header and footer for a cleaner look */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    with st.sidebar:
        st.title("Settings and Control ⚙️")

        st.markdown("### Chat Mode")
        chat_mode_selection = st.radio(
            "Select Chat Mode:",
            ("General Chat", "Document Q&A"),
            index=0 if st.session_state.current_chat_mode == "general" else 1,
            key="chat_mode_radio"
        )
        if chat_mode_selection == "General Chat":
            new_mode = "general"
        else:
            new_mode = "document_qa"

        if new_mode != st.session_state.current_chat_mode:
            st.session_state.current_chat_mode = new_mode
            st.session_state.messages = []
            if new_mode == "document_qa" and not st.session_state.vector_store_loaded:
                st.session_state.messages.append({"role": "assistant", "content": "Switched to Document Q&A. Please upload and process your PDFs first to ask questions about them."})
            elif new_mode == "general":
                st.session_state.messages.append({"role": "assistant", "content": "Switched to General Chat."})
            st.rerun()

        st.markdown("---")

        if st.button("🔄 New Chat", help="Start a new conversation", use_container_width=True):
            st.session_state.messages = []
            st.session_state.vector_store_loaded = False
            if os.path.exists("faiss_index"):
                shutil.rmtree("faiss_index")
            st.session_state.chat_session = general_chat_model.start_chat(history=[])
            st.success("Chat reset!")
            st.rerun()

        if st.session_state.current_chat_mode == "document_qa" or st.session_state.vector_store_loaded:
            st.markdown("---")
            st.title("📂 Upload Menu")
            pdf_docs = st.file_uploader("📎 Upload PDF Files", accept_multiple_files=True, key="pdf_uploader")
            if st.button("🚀 Submit & Process", use_container_width=True, key="process_pdfs_button"):
                if pdf_docs:
                    with st.spinner("Processing PDFs and building knowledge base..."):
                        documents = get_pdf_documents(pdf_docs)
                        text_chunks = get_text_chunks_from_documents(documents)
                        get_vector_store(text_chunks)
                        st.session_state.messages = []
                        st.session_state.messages.append({"role": "assistant", "content": "Knowledge base updated! What would you like to know about these documents?"})
                        st.session_state.current_chat_mode = "document_qa"
                        st.rerun()
                else:
                    st.warning("⚠️ Please upload at least one PDF file.")

        st.markdown("---")
        st.title("📘 About This App")
        st.markdown("""
        **Welcome to ChatPDF 💁‍♂️**

        🚀 Upload one or more PDF files and interact with them instantly using the power of AI.

        💬 **Ask questions** about the documents — whether it's to extract specific details, summarize long content, or clarify complex information.

        📚 **Multi-PDF support** lets you upload several documents at once and chat across all of them as if you're talking to a knowledgeable assistant.

        🎯 Built for students, researchers, professionals, or anyone who wants fast, intelligent access to PDF content without manually reading through it.
        """)

    user_prompt = st.chat_input("Ask me anything...")

    if user_prompt:
        st.session_state.messages.append({"role": "user", "content": user_prompt})
        with st.chat_message("user"):
            st.markdown(user_prompt)

        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            full_response = ""

            if st.session_state.current_chat_mode == "document_qa":
                if st.session_state.vector_store_loaded and os.path.exists("faiss_index"):
                    try:
                        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
                        new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
                        docs = new_db.similarity_search(user_prompt)

                        qa_model, qa_prompt_template = get_conversational_chain_components()

                        context_parts = []
                        for doc in docs:
                            page_info = f" (Page {doc.metadata.get('page', 'N/A')})"
                            context_parts.append(f"{doc.page_content}{page_info}")
                        context_text = "\n\n".join(context_parts)

                        formatted_prompt = qa_prompt_template.format(context=context_text, question=user_prompt)

                        document_response_stream = qa_model.stream(formatted_prompt)
                        document_response_text = ""
                        for chunk in document_response_stream:
                            document_response_text += chunk.content
                            message_placeholder.markdown(document_response_text)
                            time.sleep(0.02)
                        message_placeholder.markdown(document_response_text)

                        if "Answer is not available in the context." in document_response_text:
                            full_response += document_response_text + "\n\n"
                            full_response += "**Additional Information**\n\n"
                            message_placeholder.markdown(full_response)

                            general_chat_response_stream = st.session_state.chat_session.send_message(user_prompt, stream=True)
                            for chunk in general_chat_response_stream:
                                full_response += chunk.text
                                message_placeholder.markdown(full_response)
                                time.sleep(0.02)
                            message_placeholder.markdown(full_response)
                            ai_response_text = full_response
                        else:
                            ai_response_text = document_response_text

                    except Exception as e:
                        ai_response_text = f"An error occurred while retrieving information from PDFs: {e}\n\n"
                        ai_response_text += "Attempting General Knowledge Fallback\n\n"
                        try:
                            general_chat_response_stream = st.session_state.chat_session.send_message(user_prompt, stream=True)
                            for chunk in general_chat_response_stream:
                                ai_response_text += chunk.text
                                message_placeholder.markdown(ai_response_text)
                                time.sleep(0.02)
                            message_placeholder.markdown(ai_response_text)
                        except Exception as e_general:
                            ai_response_text += f"An error occurred during general chat fallback: {e_general}"
                else:
                    ai_response_text = "Please upload and process your PDFs in the sidebar first to use Document Q&A mode.\n\n"
                    ai_response_text = "**You are currently in General Chat Mode. Please upload and submit a file to get answers.**\n\n"
                    try:
                        general_chat_response_stream = st.session_state.chat_session.send_message(user_prompt, stream=True)
                        for chunk in general_chat_response_stream:
                            ai_response_text += chunk.text
                            message_placeholder.markdown(ai_response_text)
                            time.sleep(0.02)
                        message_placeholder.markdown(ai_response_text)
                    except Exception as e_general:
                        ai_response_text += f"An error occurred during general chat: {e_general}"
            else:
                try:
                    gemini_response = st.session_state.chat_session.send_message(user_prompt, stream=True)
                    for chunk in gemini_response:
                        full_response += chunk.text
                        message_placeholder.markdown(full_response)
                        time.sleep(0.02)
                    ai_response_text = full_response
                except Exception as e:
                    ai_response_text = f"An error occurred during general chat: {e}"

            message_placeholder.markdown(ai_response_text)
            st.session_state.messages.append({"role": "assistant", "content": ai_response_text})
            

if __name__ == "__main__":
    main()
